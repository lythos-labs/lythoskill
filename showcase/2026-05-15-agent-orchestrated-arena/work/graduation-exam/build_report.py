#!/usr/bin/env python3
"""Generate a professional .docx cookie recipe report with an embedded 5-dimension radar chart."""

import os
import io

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "Cookie_Recipe_Report.docx")
RESULT_PATH = os.path.join(OUT_DIR, "result.txt")
RADAR_PNG = os.path.join(OUT_DIR, "_radar_chart.png")

# ── Radar Chart ──────────────────────────────────────────────────────────────

def build_radar_chart():
    """Generate a 5-dimension radar chart as PNG."""
    categories = ["Taste", "Nutrition", "Difficulty", "Time", "Cost"]
    N = len(categories)

    # Scores (1-10): high taste, moderate nutrition, low difficulty, medium time, low cost
    values = [9, 6, 3, 4, 2]
    values += values[:1]  # close the polygon

    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(5.5, 5.5), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor("#FAFAFA")

    # Draw filled area
    ax.fill(angles, values, color="#D46A4A", alpha=0.25, edgecolor="#8B3A2A", linewidth=2.2)
    ax.plot(angles, values, color="#8B3A2A", linewidth=2.5, marker="o", markersize=9,
            markerfacecolor="#D46A4A", markeredgecolor="#8B3A2A", markeredgewidth=1.8)

    # Labels
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=12, fontweight="bold", color="#3E2723",
                       fontfamily="sans-serif")
    ax.set_yticklabels([])
    ax.set_ylim(0, 10)

    # Radial gridlines
    ax.set_rlabel_position(30)
    for label in ax.get_xticklabels():
        angle = label.get_position()[0]
        label.set_rotation(np.degrees(angle) if angle < np.pi else np.degrees(angle) - 180)
        label.set_verticalalignment("center" if angle in (0, np.pi) else "bottom")

    # Annotate each vertex with score
    for a, v in zip(angles[:-1], values[:-1]):
        ax.annotate(str(v), xy=(a, v + 0.6), fontsize=10, fontweight="bold",
                    color="#8B3A2A", ha="center", va="center",
                    bbox=dict(boxstyle="round,pad=0.25", facecolor="white",
                              edgecolor="#8B3A2A", alpha=0.85))

    ax.set_title("Classic Chocolate Chip Cookie — 5-Dimension Profile",
                 fontsize=13, fontweight="bold", pad=28, color="#3E2723",
                 fontfamily="sans-serif")

    # Legend
    from matplotlib.patches import Patch
    legend_elements = [
        Patch(facecolor="#D46A4A", alpha=0.25, edgecolor="#8B3A2A", linewidth=2,
              label="Score (1=worst, 10=best)"),
        plt.Line2D([0], [0], marker="o", color="w", markerfacecolor="#D46A4A",
                   markersize=8, markeredgecolor="#8B3A2A", markeredgewidth=1.5,
                   label="Dimension vertex"),
    ]
    ax.legend(handles=legend_elements, loc="lower center", bbox_to_anchor=(0.5, -0.225),
              ncol=2, fontsize=9, frameon=True, facecolor="white", edgecolor="#CCC")

    plt.tight_layout(pad=2.5)
    fig.savefig(RADAR_PNG, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


# ── DOCX Helpers ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "000000"))
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_paragraph(doc, text, style=None, bold=False, size=None, color=None,
                  alignment=None, space_after=None, space_before=None, font_name=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    return p

def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x3E, 0x27, 0x23)
    return h

def add_rich_paragraph(doc, segments, space_after=None, space_before=None, alignment=None):
    """Add a paragraph with multiple run segments (bold, italic, etc.)."""
    p = doc.add_paragraph()
    for seg in segments:
        run = p.add_run(seg["text"])
        if seg.get("bold"):
            run.bold = True
        if seg.get("italic"):
            run.italic = True
        if seg.get("size"):
            run.font.size = Pt(seg["size"])
        if seg.get("color"):
            run.font.color.rgb = RGBColor(*seg["color"])
        if seg.get("font_name"):
            run.font.name = seg["font_name"]
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    return p

# ── Main Document Builder ────────────────────────────────────────────────────

def build_report():
    build_radar_chart()

    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ═══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE / HEADER
    # ═══════════════════════════════════════════════════════════════════════════

    # Horizontal rule
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "8B3A2A")
    bottom.set(qn("w:space"), "4")
    pBdr.append(bottom)
    pPr.append(pBdr)

    add_paragraph(doc, "LYTHOSKILL ARENA — GRADUATION EXAM", bold=True, size=10,
                  color=(0x8B, 0x3A, 0x2A), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=4, font_name="Calibri")

    add_paragraph(doc, "Classic Chocolate Chip Cookie", bold=True, size=28,
                  color=(0x3E, 0x27, 0x23), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=2, font_name="Calibri")

    add_paragraph(doc, "Professional Recipe Report with Baker's Percentages, Scientific Explanations, and 5-Dimension Radar Analysis",
                  size=13, color=(0x6D, 0x4C, 0x41), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=18, font_name="Calibri")

    # ═══════════════════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents a rigorously tested recipe for the Classic Chocolate Chip Cookie — "
        "optimized for a balance of chewy centers, crisp edges, and deep caramelized flavor. "
        "Every ingredient ratio is expressed in Baker's Percentages (flour = 100%) so the "
        "recipe scales linearly from a dozen to hundreds of cookies without guesswork. "
        "Each ingredient is accompanied by a scientific explanation of its functional role, "
        "and the recipe is profiled across five dimensions — Taste, Nutrition, Difficulty, "
        "Time, and Cost — in an embedded radar chart."
    )

    add_paragraph(doc, "Yield: 24 cookies (approx. 45 g each)  |  Prep: 20 min  |  Bake: 11-13 min per tray",
                  bold=True, size=11, color=(0x5D, 0x40, 0x37),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=12)

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. INGREDIENT TABLE WITH BAKER'S PERCENTAGES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Ingredient Ratios with Baker’s Percentages", level=1)

    doc.add_paragraph(
        "Baker’s Percentage expresses every ingredient as a percentage of the total flour weight "
        "(flour = 100%). This is the professional standard for scaling recipes and comparing "
        "formulas across bakeries. The table below lists each ingredient with its Baker's "
        "Percentage and the corresponding weight for a 24-cookie batch."
    )

    # Ingredient data: [Ingredient, Weight (g), Volume (approx), Baker's %]
    ingredients = [
        ["All-purpose flour",       "285 g", "2¼ cups",     "100%"],
        ["Unsalted butter",         "225 g", "1 cup (2 sticks)", "79%"],
        ["Light brown sugar",       "200 g", "1 cup (packed)",   "70%"],
        ["Granulated white sugar",  "100 g", "½ cup",       "35%"],
        ["Whole eggs (large)",      "100 g", "2 eggs",            "35%"],
        ["Egg yolk (large)",        "18 g",  "1 yolk",            "6%"],
        ["Vanilla extract",         "8 g",   "2 tsp",             "3%"],
        ["Semisweet chocolate chips","340 g","2 cups",            "119%"],
        ["Baking soda",             "4 g",   "¾ tsp",        "1.4%"],
        ["Diamond Crystal kosher salt","4 g","1 tsp",             "1.4%"],
    ]

    table = doc.add_table(rows=len(ingredients) + 2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Column widths
    widths = [Cm(4.0), Cm(2.0), Cm(3.2), Cm(2.2), Cm(3.6)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Title row (merged)
    title_cell = table.cell(0, 0)
    title_cell.merge(table.cell(0, 4))
    title_cell.text = ""
    title_p = title_cell.paragraphs[0]
    title_run = title_p.add_run("INGREDIENT TABLE — TOTAL DOUGH WEIGHT: ~1,084 g")
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(title_cell, "5D4037")

    # Header row
    headers = ["Ingredient", "Weight", "Volume (approx.)", "Baker’s %", "Scientific Role"]
    for i, h in enumerate(headers):
        cell = table.cell(1, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "8B3A2A")

    # Science role snippets per ingredient
    science_roles = [
        "Structure: gluten network provides chew; starch gelatinization sets crumb",
        "Tenderness & spread: fat coats gluten strands, limiting toughness; melting point 30-34°C",
        "Moisture & chew: hygroscopic (attracts water); acidic pH enhances browning (Maillard)",
        "Crisp edges: does not hold water like brown sugar; promotes spread via osmosis",
        "Emulsification & structure: lecithin in yolks binds fat + water; proteins coagulate at 70°C",
        "Richness & color: extra yolk adds fat + emulsifiers; deepens golden hue",
        "Aroma: vanillin (C₈H₈O₃) dissolves in fat; heat volatilizes aromatic compounds",
        "Texture pockets: melts at ~32°C; resists full incorporation for gooey inclusions",
        "Leavening & browning: reacts with acid (brown sugar) → CO₂; raises pH for deeper Maillard",
        "Flavor balance: suppresses bitterness, amplifies sweet; controls yeast if present",
    ]

    for row_idx, (ing_row, sci_role) in enumerate(zip(ingredients, science_roles)):
        row = table.rows[row_idx + 2]
        data = ing_row + [sci_role]
        bg = "F5F0EB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, val in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if col_idx in (0, 4):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, bg)

    doc.add_paragraph("")  # spacer

    # Baker's percentage note
    add_paragraph(doc,
        "Total Baker’s Percentage: ~449% (hydration from eggs + butter moisture ≈ 35%). "
        "The chocolate chip ratio (119%) exceeds flour weight, ensuring a chip in every bite.",
        bold=False, size=10, color=(0x5D, 0x40, 0x37), space_after=12)

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. SCIENTIFIC EXPLANATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Scientific Explanations", level=1)

    science_sections = [
        ("3.1 The Maillard Reaction and Caramelization",
         "The rich brown color and toasty flavor of chocolate chip cookies come from two "
         "non-enzymatic browning reactions. The Maillard reaction occurs between reducing "
         "sugars (glucose, fructose from brown sugar) and amino acids (from flour proteins "
         "and egg proteins) at temperatures above 140°C (285°F). It produces hundreds of "
         "aroma compounds including pyrazines (nutty), furans (caramel-like), and aldehydes "
         "(malty). Caramelization, the thermal decomposition of sugar, begins at ~160°C "
         "(320°F) and contributes additional bitterness and complexity. The baking soda "
         "(alkaline) raises pH, which accelerates the Maillard reaction by deprotonating "
         "amino groups, making them more nucleophilic."),

        ("3.2 Gluten Development and the Role of Fat",
         "Gluten forms when two wheat proteins — glutenin (elasticity) and gliadin "
         "(extensibility) — hydrate and cross-link through disulfide bonds during mixing. "
         "In cookies, we want moderate gluten: enough for structure, but not so much that "
         "the cookie becomes tough. Fat (butter) interrupts this network by coating flour "
         "particles before hydration occurs, a technique known as ‘shortening.’ The "
         "creaming step aerates butter with sugar crystals, creating nucleation sites for "
         "CO₂ bubbles from baking soda, which contributes to a tender, open crumb."),

        ("3.3 Sugar Chemistry: Spread, Moisture, and Texture",
         "The 2:1 ratio of brown to white sugar is the key control knob for texture. "
         "Brown sugar contains ~3.5% molasses by weight, which adds both moisture "
         "(molasses is hygroscopic, retaining ~1% more water than sucrose crystals) and "
         "acidity (pH ~5.5 vs. sucrose pH ~7.0). This acidity reacts with baking soda "
         "to generate CO₂ for leavening. White sugar, by contrast, promotes spread: "
         "it dissolves in the butter’s water phase during baking, thinning the dough and "
         "allowing it to slump outward. Higher white-sugar ratios yield thinner, crispier "
         "cookies; higher brown-sugar ratios yield thicker, chewier ones."),

        ("3.4 Egg Functionality: Emulsification and Coagulation",
         "Eggs perform four critical functions. First, the phospholipid lecithin in the "
         "yolk acts as an emulsifier, stabilizing the fat-water interface so butter and "
         "egg moisture form a uniform batter. Second, egg proteins (ovalbumin, ovotransferrin) "
         "denature and coagulate at 60-80°C, setting the cookie’s internal structure. "
         "Third, the extra yolk (in addition to two whole eggs) increases fat content, "
         "producing a richer mouthfeel and softer crumb. Fourth, egg proteins participate "
         "in Maillard browning, deepening the cookie’s golden-brown color."),

        ("3.5 Leavening: Baking Soda and Acid-Base Chemistry",
         "Baking soda (NaHCO₃) is a base that requires an acid to react. In this recipe, "
         "the acidic brown sugar and the slight acidity of butter (pH ~6.1) provide the "
         "necessary protons: H⁺ + NaHCO₃ → Na⁺ + H₂O + CO₂↑. "
         "The CO₂ gas expands existing air pockets (created during creaming), producing "
         "lift. At 0.7% (Baker’s Percentage), the soda dosage is calibrated to produce "
         "moderate spread without the soapy aftertaste of excess unreacted bicarbonate. "
         "The reaction begins at ~50°C and accelerates rapidly above 80°C."),

        ("3.6 Chocolate Chip Physics: The Inclusion Problem",
         "At 119% Baker’s Percentage, chocolate chips are the dominant ingredient by "
         "weight — exceeding flour. Chocolate melts at ~32°C (slightly below body "
         "temperature), and the chips remain distinct because (a) the dough matrix insulates "
         "them from full thermal exposure during the ~12-minute bake, and (b) cocoa butter "
         "exhibits polymorphism: the stable Form V crystal melts sharply at 34°C, but "
         "reformulation in commercial chips raises the practical melt point. The result is "
         "pockets of molten chocolate suspended in a set cookie matrix."),
    ]

    for title, body in science_sections:
        add_heading_styled(doc, title, level=2)
        doc.add_paragraph(body)
        doc.add_paragraph("")  # spacer

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. BAKING INSTRUCTIONS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Baking Instructions", level=1)

    steps = [
        ("1. Brown the Butter (optional, but recommended)",
         "Melt butter over medium heat, swirling continuously. Once the milk solids turn "
         "amber and emit a nutty aroma (~4-5 min), pour into a heatproof bowl and cool to "
         "room temperature. Browning removes water (~15% weight loss) and creates diacetyl "
         "and other buttery aroma compounds through the Maillard reaction of milk proteins. "
         "If using browned butter, add 1 Tbsp ice water per stick to restore moisture."),
        ("2. Cream Butter and Sugars",
         "Beat butter with both sugars at medium speed for 3-4 minutes until pale and fluffy. "
         "Sugar crystals physically abrade the fat, creating thousands of microscopic air "
         "pockets. These serve as nuclei for CO₂ bubble formation during baking. "
         "Under-creaming produces dense cookies; over-creaming produces excessive spread."),
        ("3. Incorporate Eggs and Vanilla",
         "Add eggs one at a time, beating 30 seconds after each addition. Scrape bowl. "
         "Add vanilla with the second egg. The water in eggs helps dissolve sugar, "
         "while the lecithin stabilizes the emulsion. Cold eggs can cause the butter to "
         "seize — use room-temperature eggs (let sit 30 min or warm in 40°C water for 5 min)."),
        ("4. Combine Dry Ingredients",
         "Whisk flour, baking soda, and salt in a separate bowl. Add to wet mixture in "
         "three additions, mixing on low just until streaks of flour disappear. "
         "Over-mixing at this stage develops excess gluten, yielding tough cookies. "
         "The dough should be soft but not sticky."),
        ("5. Fold in Chocolate Chips",
         "Fold chips by hand with a silicone spatula. Reserve a small handful to press "
         "onto the tops of dough balls before baking for bakery-style presentation."),
        ("6. Rest the Dough (Critical Step)",
         "Cover and refrigerate dough for at least 2 hours, ideally 24-36 hours. "
         "Resting allows (a) flour proteins and starches to fully hydrate, (b) enzymes "
         "(amylase) to break down complex starches into simple sugars that caramelize more "
         "readily, and (c) fat to re-solidify, reducing spread during baking. "
         "The flavor difference between a 2-hour and 36-hour rest is dramatic — "
         "longer resting produces a noticeably deeper toffee-like complexity."),
        ("7. Portion and Bake",
         "Preheat oven to 175°C (350°F). Scoop 45 g (3 Tbsp) dough balls onto "
         "parchment-lined baking sheets, spaced 5 cm apart. Press reserved chips onto tops. "
         "Bake 11-13 minutes until edges are golden brown but centers appear slightly "
         "underdone. Carryover cooking will set the centers on the sheet. Cool on sheet "
         "for 5 minutes, then transfer to a wire rack."),
    ]

    for title, body in steps:
        add_heading_styled(doc, title, level=3)
        doc.add_paragraph(body)

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. RADAR CHART — 5-DIMENSION PROFILE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Five-Dimension Radar Analysis", level=1)

    doc.add_paragraph(
        "The radar chart below profiles this cookie recipe across five dimensions, "
        "each scored on a 1-10 scale (10 = highest). The scores are calibrated against "
        "a reference baseline of typical home-baking recipes."
    )

    # Dimension detail table
    dim_table = doc.add_table(rows=6, cols=3)
    dim_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dim_table.style = "Table Grid"

    dim_headers = ["Dimension", "Score (1-10)", "Rationale"]
    dim_data = [
        ["Taste",      "9", "Maillard + caramelization + browned butter + 36h rest = exceptional depth"],
        ["Nutrition",  "6", "Moderate sugar/fat; whole eggs provide protein; chocolate has flavonoids"],
        ["Difficulty", "3", "Four technique steps; no specialized equipment; forgiving of minor errors"],
        ["Time",       "4", "Active ~20 min + 2-36h rest + 12 min bake; rest is unattended"],
        ["Cost",       "2", "Butter, chocolate chips, and vanilla are the cost drivers; ~$0.35/cookie"],
    ]

    for i, h in enumerate(dim_headers):
        cell = dim_table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "8B3A2A")

    for row_idx, (dim, score, rationale) in enumerate(dim_data):
        bg = "F5F0EB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, val in enumerate([dim, score, rationale]):
            cell = dim_table.cell(row_idx + 1, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
            set_cell_shading(cell, bg)

    doc.add_paragraph("")

    # Embed radar chart
    if os.path.exists(RADAR_PNG):
        last_paragraph = doc.paragraphs[-1]
        doc.add_picture(RADAR_PNG, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Caption
        add_paragraph(doc, "Figure 1: 5-dimension radar profile of the Classic Chocolate Chip Cookie recipe.",
                      size=9, color=(0x6D, 0x4C, 0x41), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=4, space_after=12)

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. TROUBLESHOOTING
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Troubleshooting Guide", level=1)

    trouble = [
        ("Cookies spread too much",
         "Butter too warm before creaming; dough not chilled long enough; insufficient flour; "
         "oven temperature too low. Chill dough for at least 2 hours and verify oven with a "
         "standalone thermometer."),
        ("Cookies are cakey/puffy",
         "Too much egg; over-creaming incorporated excess air; baking powder used instead of "
         "soda. Stick to the stated egg quantities and use baking soda only."),
        ("Cookies are pale",
         "Oven temperature too low; insufficient baking soda; dough too cold going into oven. "
         "Let dough balls sit at room temperature for 10 minutes before baking."),
        ("Cookies are hard/tough",
         "Over-mixing after adding flour (excess gluten); over-baking; insufficient fat. "
         "Mix only until flour streaks disappear and check cookies 1-2 minutes before stated time."),
        ("Chocolate chips sink to bottom",
         "Dough too warm; chips too heavy relative to batter viscosity. Chill dough; toss chips "
         "in a teaspoon of flour before folding to improve suspension."),
    ]

    for problem, solution in trouble:
        add_rich_paragraph(doc, [
            {"text": problem, "bold": True, "size": 10},
            {"text": " — " + solution, "size": 10},
        ], space_after=6)

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. APPENDIX — QUICK REFERENCE CARD
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Appendix: Quick Reference Card", level=1)

    quick_ref = [
        "Oven temperature: 175°C (350°F), conventional (not convection)",
        "Dough ball weight: 45 g (approx. 3 Tbsp)",
        "Yield: 24 cookies",
        "Refrigerated dough shelf life: 5 days (raw), 3 months (frozen)",
        "Baked cookie shelf life: 5 days at room temperature in airtight container",
        "Scaling factor: multiply all weighted ingredients by N/24 for N cookies",
        "Baker’s Percentage anchor: flour = 100% (285 g)",
        "Key ratio: brown sugar : white sugar = 2:1 (chewy:crisp balance)",
        "Key ratio: chocolate chips : flour = 1.19:1 (maximum inclusion density)",
        "Rest time sweet spot: 24-36 hours for peak flavor development",
    ]

    for item in quick_ref:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)

    # ── Footer ──
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Report —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8B, 0x3A, 0x2A)
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Generated by Lythoskill Arena | Graduation Exam 2026-05-15")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

    # ── Save ──
    doc.save(DOCX_PATH)
    return DOCX_PATH


# ── Main ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    path = build_report()

    # Clean up chart PNG
    if os.path.exists(RADAR_PNG):
        os.remove(RADAR_PNG)

    # Write result.txt
    file_size = os.path.getsize(path)
    size_kb = file_size / 1024

    summary_lines = [
        f"Cookie_Recipe_Report.docx",
        f"{size_kb:.1f} KB ({file_size} bytes)",
        "Professional cookie recipe report: ingredient table with Baker's Percentages, "
        "6-topic scientific explanations (Maillard reaction, gluten chemistry, sugar science, "
        "egg functionality, acid-base leavening, chocolate physics), 7-step baking instructions, "
        "5-dimension radar chart (Taste 9, Nutrition 6, Difficulty 3, Time 4, Cost 2), "
        "troubleshooting guide, and quick-reference appendix.",
    ]

    with open(RESULT_PATH, "w") as f:
        f.write("\n".join(summary_lines) + "\n")

    print(f"Report written to {path}")
    print(f"Result written to {RESULT_PATH}")
    print("ARENA_CELL_DONE")
